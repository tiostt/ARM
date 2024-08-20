from form import orders_form

from PyQt5.QtWidgets import QMainWindow, QMessageBox
from PyQt5 import QtCore

from models import *

from docx import Document, table
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyQt5.QtWidgets import QMessageBox

import os
import re
import peewee

class Addedit_orders(QMainWindow, orders_form.Ui_AddEditForm):
    def __init__(self, parent=None, orders=None):
        super(Addedit_orders, self).__init__(parent)
        self.setupUi(self)
        self.setWindowModality(QtCore.Qt.ApplicationModal)
        self.orders = orders
        self.pattern_quantity = re.compile(r'^(\d+|\d+\.\d+)$')
        
        selected_date_1 = self.get_selected_date_1()
        self.lineEdit_startDate.setText(selected_date_1.strftime("%d.%m.%Y"))
        self.calendarWidget_1.clicked.connect(self.get_selected_date_1)
        self.lineEdit_startDate.setReadOnly(True)
        
        selected_date_2 = self.get_selected_date_2()
        self.lineEdit_Date.setText(selected_date_2.strftime("%d.%m.%Y"))
        self.calendarWidget_2.clicked.connect(self.get_selected_date_2)
        self.lineEdit_Date.setReadOnly(True)

        self.loaddata_combobox()

        self.OkBtn.setDefault(True)
        self.CancelBtn.setDefault(True)

        self.OkBtn.setFocusPolicy(QtCore.Qt.StrongFocus)
        self.CancelBtn.setFocusPolicy(QtCore.Qt.StrongFocus)

        if self.orders:
            pass
        else:
            self.OkBtn.clicked.connect(self.create_orders)

        self.CancelBtn.clicked.connect(self.close)
        
    def closeEvent(self, event):
        if self.parent() is not None:
            self.parent().loaddata_order()
        super(Addedit_orders, self).closeEvent(event) 

    def get_selected_date_1(self):
        selected_date = self.calendarWidget_1.selectedDate()
        if self.lineEdit_startDate.text() == "":
            return selected_date.toPyDate()
        else:
            selected_date = selected_date.toPyDate()
            self.lineEdit_startDate.setText(selected_date.strftime("%d.%m.%Y"))
    
    def get_selected_date_2(self):
        selected_date = self.calendarWidget_2.selectedDate()
        if self.lineEdit_Date.text() == "":
            return selected_date.toPyDate()
        else:
            selected_date = selected_date.toPyDate()
            self.lineEdit_Date.setText(selected_date.strftime("%d.%m.%Y"))

    def validator(self):
        if self.lineEdit_name.text() == "":
            QMessageBox.critical(self, 'Ошибка', 'Вы не заполнили поле "Номер заказа"!')
            return False
        elif self.comboBox_clients.currentText() == "":
            QMessageBox.critical(self, 'Ошибка', 'Вы не заполнили поле "Траспортное средство" или сама таблица пуста')
            return False
        elif self.comboBox_product.currentText() == "":
            QMessageBox.critical(self, 'Ошибка', 'Вы не заполнили поле "Поставщик" или сама таблица пуста')
            return False
        elif self.comboBox_venicle.currentText() == "":
            QMessageBox.critical(self, 'Ошибка', 'Вы не заполнили поле "Продукт" или сама таблица пуста')
            return False
        elif not self.pattern_quantity.search(self.lineEdit_quantity.text()) or float(self.lineEdit_quantity.text()) == 0:
            QMessageBox.critical(self, 'Ошибка', 'Вы не корректно заполнили поле "Количество"!')
            return False
        else:
            return True

    def loaddata_combobox(self):
        self.comboBox_venicle.clear()
        self.comboBox_clients.clear()
        self.comboBox_product.clear()
        data_venicle = [x for x in Venicle.select()]
        data_clients = [x for x in Clients.select()]
        data_product = [x for x in Products.select()]
        
        for venicle in data_venicle:
            self.comboBox_venicle.addItem(venicle.model, venicle.id)
        for client in data_clients:
            client_name = f"{client.firstName} {client.lastName} {client.middleName}"
            self.comboBox_clients.addItem(client_name, client.id)
        for product in data_product:
            product_name = f"{product.name}|{product.suppliers.name}"
            self.comboBox_product.addItem(product_name, product.id)

    def create_orders(self):
        if self.validator():
            name = self.lineEdit_name.text()
            clients = self.comboBox_clients.currentData()
            product = self.comboBox_product.currentData()
            venicle = self.comboBox_venicle.currentData()
            startDate = self.lineEdit_startDate.text()
            orderdate = self.lineEdit_Date.text()

            product_quanity = Products.get(id=product)
            if float(self.lineEdit_quantity.text()) > float(product_quanity.quantity):
                QMessageBox.critical(self, 'Ошибка', 'Недостаточное количество товара на складе')
            else:
                quanity = self.lineEdit_quantity.text()

                product_price = Products.get(id=product)
                Total_price = float(quanity) * float(product_price.price)

                try:
                    Orders(
                        name=name,
                        startDate=startDate,
                        orderdate=orderdate,
                        quanity=quanity,
                        Total_price=Total_price,
                        clients=clients,
                        product=product,
                        venicle=venicle,
                    ).save()
                    self.close()
                except peewee.IntegrityError:
                    QMessageBox.critical(self, 'Ошибка', 'Заказ с таким именем уже существует!')
                except peewee.DataError:
                    QMessageBox.critical(self, 'Ошибка', 'Цена не может превышать 100 млн рублей')

    def edit_orders(self, orders):
        self.lineEdit_name.setText(orders.name)

        index = self.comboBox_clients.findData(orders.clients.id)
        self.comboBox_clients.setCurrentIndex(index)

        index = self.comboBox_product.findData(orders.product.id)
        self.comboBox_product.setCurrentIndex(index)

        index = self.comboBox_venicle.findData(orders.venicle.id)
        self.comboBox_venicle.setCurrentIndex(index)

        self.lineEdit_quantity.setText(str(orders.quanity))
        self.lineEdit_startDate.setText(str(orders.startDate.strftime('%d.%m.%Y')))
        self.lineEdit_Date.setText(str(orders.orderdate.strftime('%d.%m.%Y')))
        self.calendarWidget_1.setSelectedDate(orders.startDate)
        self.calendarWidget_2.setSelectedDate(orders.orderdate)

        self.OkBtn.clicked.connect(lambda: self.update_orders(orders))

    def update_orders(self, orders):
        if self.validator():
            update = Orders.get(Orders.id == orders.id)
            update.name = self.lineEdit_name.text()
            update.clients = self.comboBox_clients.currentData()
            update.product = self.comboBox_product.currentData()
            update.venicle = self.comboBox_venicle.currentData()
            update.startDate  = self.lineEdit_startDate.text()
            update.orderdate  = self.lineEdit_Date.text()
            
            if float(self.lineEdit_quantity.text()) > float(orders.product.quantity):
                QMessageBox.critical(self, 'Ошибка', 'Недостаточное количество товара на складе')
            else:
                update.quanity = self.lineEdit_quantity.text()

                update.Total_price = float(self.lineEdit_quantity.text()) * float(orders.product.price)

                try:
                    update.save()
                    self.close()
                except peewee.IntegrityError:
                    QMessageBox.critical(self, 'Ошибка', 'Заказ с таким именем уже существует!')
                except peewee.DataError:
                    QMessageBox.critical(self, 'Ошибка', 'Цена не может превышать 100 млн рублей')

    def delete_orders(self, orders):
        msg = QMessageBox()
        msg.setWindowTitle("Подтверждение действия")
        msg.setText(f"Вы уверены, что заказ '{orders.name}' аннулирован?")
        msg.setIcon(QMessageBox.Warning)
        msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)

        # Задаем русские названия кнопкам
        msg.button(QMessageBox.Yes).setText('Да')
        msg.button(QMessageBox.No).setText('Нет')

        confirm = msg.exec_()
        if confirm == QMessageBox.Yes:
            try:
                orders.delete_instance()
            except peewee.IntegrityError:
                QMessageBox.critical(self, 'Ошибка', f'На запись "{orders.name}" есть ссылка в другой таблице')
            
            if self.parent() is not None:
                self.parent().loaddata_order()  # Обновляем данные в таблице
   
    def win_orders(self, orders):
        msg = QMessageBox()
        msg.setWindowTitle("Подтверждение действия")
        msg.setText(f"Вы уверены, что заказ '{orders.name}' выполнен?")
        msg.setIcon(QMessageBox.Warning)
        msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)

        # Задаем русские названия кнопкам
        msg.button(QMessageBox.Yes).setText('Да')
        msg.button(QMessageBox.No).setText('Нет')

        confirm = msg.exec_()
        if confirm == QMessageBox.Yes:
            product = None
            try:
                product = Products.get(name=orders.product.name, suppliers=orders.product.suppliers)
                
                product_quantity = product.quantity - orders.quanity

                if product_quantity < 0:
                    QMessageBox.critical(self, 'Ошибка', 'Недостаточное количество товара на складе')
                else:
                    if product_quantity == 0:
                        QMessageBox.warning(self, 'Предупреждение', 'Товар на складе закончился и был удален из таблицы')
                        orders.delete_instance()
                        product.delete_instance()
                    else:
                        product.quantity = product_quantity
                        product.save()
                        self.create_otchet(orders)
                        orders.delete_instance()


            except Products.DoesNotExist:
                QMessageBox.critical(self, 'Ошибка', 'Такого товара не существует')

            if self.parent() is not None:
                self.parent().loaddata_order()
                self.parent().loaddata_Product()

    def create_otchet(self, order):
        # Create a new Word document
        document = Document()

        # Add a title and subtitle to the document
        document.add_heading('Отчет о заказе', 0)
        document.add_paragraph('Информация о заказе:', style='Subtitle')

        # Create a table to store the order data
        table = document.add_table(rows=1, cols=2, style='Table Grid')

        # Add the table headers
        hdr_cells = table.add_row().cells
        hdr_cells[0].text = 'Свойство'
        hdr_cells[1].text = 'Значение'

        # Add the order data to the table
        row_cells = table.add_row().cells
        row_cells[0].text = 'ID заказа'
        row_cells[1].text = str(order.id)

        row_cells = table.add_row().cells
        row_cells[0].text = 'Название заказа'
        row_cells[1].text = order.name

        row_cells = table.add_row().cells
        row_cells[0].text = 'Дата начала заказа'
        row_cells[1].text = order.startDate.strftime('%d/%m/%Y')

        row_cells = table.add_row().cells
        row_cells[0].text = 'Дата заказа'
        row_cells[1].text = order.orderdate.strftime('%d/%m/%Y')

        row_cells = table.add_row().cells
        row_cells[0].text = 'Количество заказанного товара'
        row_cells[1].text = str(order.quanity)

        row_cells = table.add_row().cells
        row_cells[0].text = 'Общая стоимость заказа'
        row_cells[1].text = str(order.Total_price)

        row_cells = table.add_row().cells
        row_cells[0].text = 'Клиент'
        row_cells[1].text = f'{order.clients.firstName} {order.clients.lastName}'

        row_cells = table.add_row().cells
        row_cells[0].text = 'Продукт'
        row_cells[1].text = order.product.name

        row_cells = table.add_row().cells
        row_cells[0].text = 'Транспортное средство'
        row_cells[1].text = order.venicle.model

        # Set the table width to fit the page
        table.width = Inches(8.5)

        # Autofit the table columns to the content
        for col in table.columns:
            col.width = Inches(1.5)

        # Add a summary paragraph
        paragraph = document.add_paragraph()
        run = paragraph.add_run('Данные актуальны на момент формирования отчета.')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 153)  # Синий цвет
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание по правому краю


        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")

        if not os.path.exists(desktop_path):
            QMessageBox.critical(self, 'Ошибка', 'Папка "Рабочий стол" не найдена.')
            return

        # Проверка на существование файла
        file_name = f'отчет_заказ_{order.name}.docx'
        file_path = os.path.join(desktop_path, file_name)
        if os.path.exists(file_path):
            # Если файл существует, генерируем уникальное имя
            i = 1
            while os.path.exists(os.path.join(desktop_path, f'отчет_заказ_{order.name}_{i}.docx')):
                i += 1
            file_path = os.path.join(desktop_path, f'отчет_заказ_{order.name}_{i}.docx')

        # Сохранение документа
        document.save(file_path)

        QMessageBox.information(self, '', f'Отчет о записи "{order.name}" успешно сформирован')
