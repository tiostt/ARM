import os
from form import delivery_form

from PyQt5.QtWidgets import QMainWindow, QMessageBox
from PyQt5 import QtCore

from models import Delivery, Products, Suppliers, Venicle

from docx import Document, table
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import re
import peewee

class Addedit_delivery(QMainWindow, delivery_form.Ui_AddEditForm):
    def __init__(self, parent=None, delivery=None):
        super(Addedit_delivery, self).__init__(parent)
        self.setupUi(self)
        self.setWindowModality(QtCore.Qt.ApplicationModal)
        self.delivery = delivery
        self.pattern_quantity = re.compile(r'^(\d+|\d+\.\d+)$')
        
        selected_date_1 = self.get_selected_date_1()
        self.lineEdit_startDate.setText(selected_date_1.strftime("%d.%m.%Y"))
        self.calendarWidget_1.clicked.connect(self.get_selected_date_1)
        self.lineEdit_startDate.setReadOnly(True)
        
        selected_date_2 = self.get_selected_date_2()
        self.lineEdit_deliveryDate.setText(selected_date_2.strftime("%d.%m.%Y"))
        self.calendarWidget_2.clicked.connect(self.get_selected_date_2)
        self.lineEdit_deliveryDate.setReadOnly(True)

        self.loaddata_combobox()

        self.OkBtn.setDefault(True)
        self.CancelBtn.setDefault(True)
        
        self.OkBtn.setFocusPolicy(QtCore.Qt.StrongFocus)
        self.CancelBtn.setFocusPolicy(QtCore.Qt.StrongFocus)

        if self.delivery:
            pass
        else:
            self.OkBtn.clicked.connect(self.create_delivery)

        self.CancelBtn.clicked.connect(self.close)
        
    def closeEvent(self, event):
        if self.parent() is not None:
            self.parent().loaddata_delivery()
        super(Addedit_delivery, self).closeEvent(event) 

    def get_selected_date_1(self):
        selected_date = self.calendarWidget_1.selectedDate()
        if self.lineEdit_startDate.text() == "":
            return selected_date.toPyDate()
        else:
            selected_date = selected_date.toPyDate()
            self.lineEdit_startDate.setText(selected_date.strftime("%d.%m.%Y"))
    
    def get_selected_date_2(self):
        selected_date = self.calendarWidget_2.selectedDate()
        if self.lineEdit_deliveryDate.text() == "":
            return selected_date.toPyDate()
        else:
            selected_date = selected_date.toPyDate()
            self.lineEdit_deliveryDate.setText(selected_date.strftime("%d.%m.%Y"))

    def validator(self):
        if self.lineEdit_name.text() == "" or len(self.lineEdit_name.text()) <= 2:
            QMessageBox.critical(self, 'Ошибка', 'Вы не заполнили поле "Номер поставки"!')
            return False
        elif self.comboBox_venicle.currentText() == "":
            QMessageBox.critical(self, 'Ошибка', 'Вы не заполнили поле "Траспортное средство" или сама таблица пуста')
            return False
        elif self.comboBox_supplier.currentText() == "":
            QMessageBox.critical(self, 'Ошибка', 'Вы не заполнили поле "Поставщик" или сама таблица пуста')
            return False
        elif self.lineEdit_product.text() == "" or len(self.lineEdit_product.text()) <= 2:
            QMessageBox.critical(self, 'Ошибка', 'Вы не заполнили поле "Товар"!')
            return False
        elif not self.pattern_quantity.search(self.lineEdit_quantity.text()) or self.lineEdit_quantity.text() == 0:
            QMessageBox.critical(self, 'Ошибка', 'Вы не корректно заполнили поле "Количество"!')
            return False
        elif not self.pattern_quantity.search(self.lineEdit_price.text()) or self.lineEdit_quantity.text() == 0:
            QMessageBox.critical(self, 'Ошибка', 'Вы не корректно заполнили поле "Общая стоимость"!')
            return False
        else:
            return True

    def loaddata_combobox(self):
        self.comboBox_venicle.clear()
        self.comboBox_supplier.clear()
        data_venicle = [x for x in Venicle.select()]
        data_supplier = [x for x in Suppliers.select()]
        
        for venicle in data_venicle:
            self.comboBox_venicle.addItem(venicle.model, venicle.id)
        for supplier in data_supplier:
            self.comboBox_supplier.addItem(supplier.name, supplier.id)

    def create_delivery(self):
        if self.validator():
            name = self.lineEdit_name.text()
            venicle = self.comboBox_venicle.currentData()
            supplier = self.comboBox_supplier.currentData()
            product = self.lineEdit_product.text()
            startDate = self.lineEdit_startDate.text()
            deliveryDate = self.lineEdit_deliveryDate.text()
            quantity = float(self.lineEdit_quantity.text())
            price = float(self.lineEdit_price.text())
            Total_price = str(price * quantity)

            try:
                Delivery(
                        name=name,
                        venicle=venicle,
                        supplier=supplier,
                        product=product,
                        startDate=startDate,
                        deliveryDate=deliveryDate,
                        quantity=quantity,
                        price=price,
                        Total_price=Total_price,
                    ).save()
                self.close() 
            except peewee.IntegrityError:
                QMessageBox.critical(self, 'Ошибка', 'Поставка с таким именем уже существует!')
            except peewee.DataError:
                QMessageBox.critical(self, 'Ошибка', 'Цена не может превышать 100 млн рублей')

    def edit_delivery(self, delivery):
        self.lineEdit_name.setText(delivery.name)

        index = self.comboBox_venicle.findData(delivery.venicle.id)
        self.comboBox_venicle.setCurrentIndex(index)

        index = self.comboBox_supplier.findData(delivery.supplier.id)
        self.comboBox_supplier.setCurrentIndex(index)


        self.lineEdit_quantity.setText(str(delivery.quantity))
        self.lineEdit_price.setText(str(delivery.price ))
        self.lineEdit_product.setText(str(delivery.product ))
        self.lineEdit_startDate.setText(str(delivery.startDate.strftime('%d.%m.%Y')))
        self.lineEdit_deliveryDate.setText(str(delivery.deliveryDate.strftime('%d.%m.%Y')))
        self.calendarWidget_1.setSelectedDate(delivery.startDate)
        self.calendarWidget_2.setSelectedDate(delivery.deliveryDate)

        self.OkBtn.clicked.connect(lambda: self.update_delivery(delivery))

    def update_delivery(self, delivery):
        if self.validator():
            update = Delivery.get(Delivery.id == delivery.id)
            update.name = self.lineEdit_name.text()
            update.product = self.lineEdit_product.text()
            update.venicle = self.comboBox_venicle.currentData()
            update.supplier = self.comboBox_supplier.currentData()
            update.startDate  = self.lineEdit_startDate.text()
            update.deliveryDate  = self.lineEdit_deliveryDate.text()
            update.quantity = self.lineEdit_quantity.text()
            update.price = self.lineEdit_price.text()
            update.Total_price = float(self.lineEdit_price.text()) * float(self.lineEdit_quantity.text())
            
            try:
                update.save()
                self.close()
            except peewee.IntegrityError:
                QMessageBox.critical(self, 'Ошибка', 'Поставка с таким именем уже существует!')
            except peewee.DataError:
                QMessageBox.critical(self, 'Ошибка', 'Цена не может превышать 100 млн рублей')

    def delete_delivery(self, delivery):
        msg = QMessageBox()
        msg.setWindowTitle("Подтверждение действия")
        msg.setText(f"Вы уверены, что поставка '{delivery.name}' аннулирована?")
        msg.setIcon(QMessageBox.Warning)
        msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)

        # Задаем русские названия кнопкам
        msg.button(QMessageBox.Yes).setText('Да')
        msg.button(QMessageBox.No).setText('Нет')

        confirm = msg.exec_()
        if confirm == QMessageBox.Yes:
            try:
                delivery.delete_instance()
            except peewee.IntegrityError:
                QMessageBox.critical(self, 'Ошибка', f'На запись "{delivery.name}" есть ссылка в другой таблице')
            if self.parent() is not None:
                self.parent().loaddata_delivery()  # Обновляем данные в таблице

    def win_delivery(self, delivery):
        msg = QMessageBox()
        msg.setWindowTitle("Подтверждение действия")
        msg.setText(f"Вы уверены, что поставка '{delivery.name}' выполнена?")
        msg.setIcon(QMessageBox.Warning)
        msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)

        # Задаем русские названия кнопкам
        msg.button(QMessageBox.Yes).setText('Да')
        msg.button(QMessageBox.No).setText('Нет')

        confirm = msg.exec_()
        if confirm == QMessageBox.Yes:

            try:
                product = Products.get(name=delivery.product, suppliers=delivery.supplier)
                product.quantity += delivery.quantity
                product.price=delivery.price

                product.save()

            except Products.DoesNotExist:
                Products(
                    name=delivery.product,
                    suppliers=delivery.supplier.id,
                    quantity=delivery.quantity,
                    price=delivery.price,
                ).save()
            self.create_otchet(delivery)
            delivery.delete_instance()

            if self.parent() is not None:
                self.parent().loaddata_delivery()
                self.parent().loaddata_Product()
    
    def create_otchet(self, delivery):
        # Create a new Word document
        document = Document()

        # Add a title and subtitle to the document
        document.add_heading('Отчет о поставки', 0)
        document.add_paragraph('Информация о поставке:', style='Subtitle')

        # Create a table to store the delivery data
        table = document.add_table(rows=1, cols=2, style='Table Grid') # Исправленная строка

        # Add the table headers
        hdr_cells = table.add_row().cells
        hdr_cells[0].text = 'Свойство'
        hdr_cells[1].text = 'Значение'

        # Add the delivery data to the table
        row_cells = table.add_row().cells
        row_cells[0].text = 'ID поставки'
        row_cells[1].text = str(delivery.id)

        row_cells = table.add_row().cells
        row_cells[0].text = 'Название поставки'
        row_cells[1].text = delivery.name

        row_cells = table.add_row().cells
        row_cells[0].text = 'Количество доставленного товара'
        row_cells[1].text = str(delivery.quantity)

        row_cells = table.add_row().cells
        row_cells[0].text = 'Поставщик'
        row_cells[1].text = delivery.supplier.name

        row_cells = table.add_row().cells
        row_cells[0].text = 'Продукт'
        row_cells[1].text = delivery.product

        row_cells = table.add_row().cells
        row_cells[0].text = 'Дата доставки'
        row_cells[1].text = delivery.deliveryDate.strftime('%d/%m/%Y')

        row_cells = table.add_row().cells
        row_cells[0].text = 'Цена за единицу товара'
        row_cells[1].text = str(delivery.price)

        row_cells = table.add_row().cells
        row_cells[0].text = 'Общая стоимость поставки'
        row_cells[1].text = str(delivery.Total_price)

        # Set the table width to fit the page
        table.width = Inches(8.5)

        # Autofit the table columns to the content
        for col in table.columns:
            col.width = Inches(1.5)

        # Add a summary paragraph
        paragraph = document.add_paragraph()
        run = paragraph.add_run('Данные актуальны на момент формирования отчета.')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 153) # Синий цвет
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Выравнивание по правому краю


        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")

        if not os.path.exists(desktop_path):
            QMessageBox.critical(self, 'Ошибка', 'Папка "Рабочий стол" не найдена.')
            return

        # Проверка на существование файла
        file_name = f'отчет_поставка_{delivery.name}.docx'
        file_path = os.path.join(desktop_path, file_name)
        if os.path.exists(file_path):
            # Если файл существует, генерируем уникальное имя
            i = 1
            while os.path.exists(os.path.join(desktop_path, f'отчет_поставка_{delivery.name}_{i}.docx')):
                i += 1
            file_path = os.path.join(desktop_path, f'отчет_поставка_{delivery.name}_{i}.docx')

        # Сохранение документа
        document.save(file_path)

        QMessageBox.information(self, '', f'Отчет о записи "{delivery.name}" успешно сформирован')

